import os
import logging
from docx import Document
from docx.oxml.ns import qn
from dotenv import load_dotenv
import base64
import sys
from io import BytesIO
from PIL import Image
from wand.image import Image as WandImage
import subprocess
from pathlib import Path

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '../../../')))

from src.core.image_describer.base_image_describer import BaseImageDescriber
from src.core.llm_client import LLMClientFactory
from src.core.image_describer.prompts.docx_prompts import system_prompt, prompt_template

# --- Cấu hình logging ---
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.StreamHandler()]
)

ALLOWED_EXT = {"png", "jpg", "jpeg", "emf", "x-emf"}
MAX_SIZE = (800, 800)


class DocxImageDescriber(BaseImageDescriber):
    """
    Trình mô tả ảnh cho tài liệu Word (file .doc và .docx).
    
    Hỗ trợ:
    - File .docx trực tiếp 
    - File .doc bằng cách chuyển đổi sang .docx qua LibreOffice
    - Chuyển đổi định dạng ảnh EMF/X-EMF
    - Thay đổi kích thước và tối ưu hóa ảnh
    - Mô tả ảnh có ngữ cảnh sử dụng LLM
    """
    def __init__(self, llm_client):
        super().__init__(llm_client)
        self.factory = LLMClientFactory()
        logging.info("📋 DocxImageDescriber đã sẵn sàng (hỗ trợ .doc và .docx)")
    def replace_images_with_description(self):
        """Hiện tại chưa cần implement, giữ nguyên để tránh lỗi abstract class."""
        logging.info("⚠️ replace_images_with_description chưa được triển khai.")
        return None
    
    def _convert_doc_to_docx(self, file_path):
        """Chuyển đổi file .doc sang .docx sử dụng LibreOffice."""
        docx_path = str(Path(file_path).with_suffix('.docx'))
        try:
            logging.info(f"🔄 Converting {file_path} to {docx_path}...")
            # Trên Windows, file thực thi LibreOffice là soffice
            subprocess.run([
                "soffice",
                "--headless",
                "--convert-to", "docx",
                "--outdir", str(Path(file_path).parent),
                file_path
            ], check=True)
            logging.info(f"✅ Successfully converted to {docx_path}")
            return docx_path
        except subprocess.CalledProcessError as e:
            logging.error(f"❌ Failed to convert .doc to .docx: {e}")
            raise RuntimeError("Failed to convert .doc to .docx using LibreOffice. Make sure LibreOffice is installed.")
        except FileNotFoundError:
            logging.error("❌ LibreOffice not found. Please install LibreOffice.")
            raise RuntimeError("LibreOffice not found. Please install LibreOffice to convert .doc files.")

    def _resize_if_needed(self, image_data, ext):
        """Resize ảnh nếu quá to, sau đó trả về bytes đã resize."""
        try:
            img = Image.open(BytesIO(image_data))
            if img.width > MAX_SIZE[0] or img.height > MAX_SIZE[1]:
                img.thumbnail(MAX_SIZE)
                buffer = BytesIO()
                format_name = "PNG" if ext.lower() == "png" else "JPEG"
                img.save(buffer, format=format_name)
                logging.info(f" Ảnh được resize về tối đa {MAX_SIZE}")
                return buffer.getvalue()
            return image_data
        except Exception as e:
            logging.warning(f" Không thể resize ảnh: {e}")
            return image_data

    def _convert_emf_to_png(self, image_data):
#         Tải ImageMagick bản hỗ trợ EMF/WMF + MagickWand

#         Vào: https://imagemagick.org/script/download.php#windows

#         Tải bản "ImageMagick-<version>-Q16-HDRI-dll.exe" (không chọn portable).
        """Chuyển EMF hoặc X-EMF → PNG bằng Wand."""
        try:
            with WandImage(blob=image_data, format="emf") as img:
                img.format = "png"
                return img.make_blob()
        except Exception as e:
            logging.error(f"Không thể convert EMF/X-EMF sang PNG: {e}")
            return None

    def extract_images_and_contexts(self, run, image_index):
        """Trích xuất ảnh dạng base64 từ run, xử lý cả EMF và X-EMF."""
        images_base64 = []
        try:
            drawings = run._element.findall(
                ".//w:drawing",
                namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            )
            for drawing in drawings:
                blip = drawing.find(
                    ".//a:blip",
                    namespaces={"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
                )
                if blip is None:
                    continue
                rId = blip.get(qn('r:embed'))
                image_part = run.part.related_parts[rId]
                image_data = image_part.blob
                ext = image_part.content_type.split("/")[-1].lower()

                if ext not in ALLOWED_EXT:
                    logging.info(f"Bỏ qua ảnh có định dạng không hỗ trợ: {ext}")
                    continue

                if ext in ("emf", "x-emf"):
                    converted = self._convert_emf_to_png(image_data)
                    if not converted:
                        continue
                    image_data = converted
                    ext = "png"

                image_data = self._resize_if_needed(image_data, ext)

                base64_image = base64.b64encode(image_data).decode("utf-8")
                images_base64.append((base64_image, ext))
                image_index += 1
        except Exception as e:
            logging.error(f"Lỗi khi trích xuất ảnh: {e}")
        return images_base64, image_index

    def generate_descriptions(self, base64_image, ext, before_text, after_text):
        """Sinh mô tả từ ảnh base64 và context."""
        try:
            image_url = f"data:image/{ext};base64,{base64_image}"
            prompt_text = prompt_template.format(before_text=before_text, after_text=after_text)

            message = self.llm["client"].chat.completions.create(
                model=self.llm["model"],
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": [
                        {"type": "text", "text": prompt_text},
                        {"type": "image_url", "image_url": {"url": image_url}}
                    ]}
                ],
            )
            description = message.choices[0].message.content.strip()
            logging.info("Đã sinh mô tả thành công")
            logging.info(f"Mô tả: {description}")
            return description
        except Exception as e:
            logging.error(f"Lỗi khi mô tả ảnh: {e}")
            return "[Không thể sinh mô tả: lỗi hệ thống]"

    def get_nearest_text(self, paragraphs, start_idx, direction):
        """Tìm paragraph gần nhất có chữ theo hướng chỉ định."""
        idx = start_idx + direction
        total = len(paragraphs)

        while 0 <= idx < total:
            text = paragraphs[idx].text.strip()
            if text:
                return text
            idx += direction
        return ""

    def run(self, file_path, output_path=None):
        """
        Xử lý file doc/docx, thay ảnh bằng mô tả và giữ nguyên format.
        """
        # Kiểm tra xem file có tồn tại không
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File không tồn tại: {file_path}")
        
        # Kiểm tra phần mở rộng file và chuyển đổi nếu cần
        ext = os.path.splitext(file_path)[1].lower()
        input_path = file_path
        
        logging.info(f"🎯 Bắt đầu xử lý file {ext.upper()}: {file_path}")
        
        # Xử lý file .doc bằng cách chuyển đổi sang .docx trước
        if ext == ".doc":
            input_path = self._convert_doc_to_docx(file_path)
            # Nếu không chỉ định đường dẫn output, sử dụng thư mục của file docx đã chuyển đổi
            if output_path is None:
                output_path = input_path.replace(".docx", "_described.docx")
        elif ext == ".docx":
            # Đối với file .docx, sử dụng đường dẫn gốc
            if output_path is None:
                output_path = file_path.replace(".docx", "_described.docx")
        else:
            raise ValueError("❌ Chỉ hỗ trợ file .doc hoặc .docx. File được cung cấp có định dạng: " + ext)
        
        logging.info(f"📖 Đang xử lý file: {input_path}")
        logging.info(f"💾 Output sẽ được lưu tại: {output_path}")
        doc = Document(input_path)
        image_index = 1
        
        for idx, para in enumerate(doc.paragraphs):
            runs_to_replace = []
            for run in para.runs:
                if 'graphic' in run._element.xml:
                    runs_to_replace.append(run)

            if runs_to_replace:
                # Lấy context chỉ một lần cho cả đoạn
                context_before = self.get_nearest_text(doc.paragraphs, idx, direction=-1)
                context_after = self.get_nearest_text(doc.paragraphs, idx, direction=1)
                
                for run in runs_to_replace:
                    images_base64, image_index = self.extract_images_and_contexts(run, image_index)
                    for base64_image, ext in images_base64:
                        description = self.generate_descriptions(base64_image, ext, context_before, context_after)
                        
                        # Thay thế nội dung của run chứa ảnh
                        run.text = "[IMAGE_DESCRIPTION]:" + description
                        
                        # Xóa bỏ phần graphic từ run
                        for child in run._element:
                            if qn("w:drawing") in child.tag:
                                run._element.remove(child)

        doc.save(output_path)
        logging.info(f"✅ File đã lưu: {output_path}")
        
        # Dọn dẹp file tạm đã chuyển đổi nếu file gốc là .doc
        if ext == ".doc" and input_path != file_path:
            try:
                os.remove(input_path)
                logging.info(f"🗑️ Đã dọn dẹp file tạm: {input_path}")
            except Exception as e:
                logging.warning(f"⚠️ Không thể dọn dẹp file tạm {input_path}: {e}")

        return output_path

# --- Chương trình chính ---
if __name__ == "__main__":
    load_dotenv()
    llm_factory = LLMClientFactory()
    gpt_client = llm_factory.get_client("gpt-4.1")
    docx = DocxImageDescriber(llm_client=gpt_client)

    # Ví dụ sử dụng - hoạt động với cả file .doc và .docx
    input_path = "data/sample_documents/elpasoamiandmdmsrfpv19_10012019-final.docx" 
    output_path = "data/sample_documents/elpasoamiandmdmsrfpv19_10012019-final_described.docx"
    docx.run(input_path, output_path)