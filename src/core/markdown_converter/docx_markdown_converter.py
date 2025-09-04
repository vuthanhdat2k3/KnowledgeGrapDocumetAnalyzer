import os
import logging
import re
import subprocess
import time
from pathlib import Path
from tqdm import tqdm
from docx import Document

from src.core.llm_client import LLMClientFactory
from src.core.markdown_converter.prompts.docx_prompts import (
    DOCX_TO_MARKDOWN_SYSTEM_PROMPT,
    DOCX_TO_MARKDOWN_DETAILED_PROMPT,
)

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

class DocxMarkdownConverter:
    def __init__(self, model_name: str = "gemini-2.5-flash", max_token: int = 10000):
        self.factory = LLMClientFactory()
        self.model_key = model_name
        self.max_token = max_token

    def _normalize_filename(self, filename: str) -> str:
        """Chuẩn hóa tên file (Windows-safe)."""
        invalid_chars = r'[<>:"/\\|?*]'
        filename = re.sub(invalid_chars, "-", filename)
        filename = re.sub(r"\s+", " ", filename)
        return filename.strip()

    def _extract_text_from_docx(self, file_path: str) -> str:
        """Trích xuất nội dung DOCX theo đúng thứ tự (paragraph + table)."""
        try:
            doc = Document(file_path)
            text_content = []

            for element in doc.element.body:
                if element.tag.endswith("p"):
                    for para in doc.paragraphs:
                        if para._element == element:
                            if para.text.strip():
                                text_content.append(para.text.strip())
                            break

                elif element.tag.endswith("tbl"):
                    for table in doc.tables:
                        if table._element == element:
                            table_text = []
                            for row in table.rows:
                                row_text = [cell.text.strip() for cell in row.cells]
                                table_text.append(" | ".join(row_text))
                            text_content.append("\n".join(table_text))
                            break

            return "\n\n".join(text_content)
        except Exception as e:
            logging.error(f"Error extracting text from DOCX: {e}")
            raise

    def _split_text(self, text: str) -> list:
        """Chia văn bản thành các chunk nhỏ để gửi LLM."""
        paragraphs = text.split("\n\n")
        chunks, current = [], ""

        for para in paragraphs:
            if len(current) + len(para) + 2 <= self.max_token:
                current += para + "\n\n"
            else:
                if current.strip():
                    chunks.append(current.strip())
                current = para + "\n\n"

        if current.strip():
            chunks.append(current.strip())
        return chunks

    def _build_prompt(self, content: str) -> list:
        """Xây dựng prompt cho Gemini theo format system + user."""
        return [
            {"role": "user", "parts": [{"text": DOCX_TO_MARKDOWN_SYSTEM_PROMPT}]},
            {"role": "user", "parts": [{"text": DOCX_TO_MARKDOWN_DETAILED_PROMPT.format(content=content)}]},
        ]

    def _extract_response_text(self, response, idx: int) -> str:
        """Trích xuất văn bản từ response của Gemini/OpenAI, fallback khi rỗng."""
        texts = []

        # Ưu tiên Gemini candidates
        if hasattr(response, "candidates"):
            for cand in getattr(response, "candidates", []):
                if getattr(cand, "content", None):
                    for part in getattr(cand.content, "parts", []):
                        if hasattr(part, "text") and part.text:
                            texts.append(part.text)

        # Fallback sang response.text
        if not texts and hasattr(response, "text"):
            if response.text:
                texts.append(response.text)

        if not texts:
            logging.warning(f"Chunk {idx+1} trả về rỗng → dùng placeholder")
            return f"[Không thể tạo nội dung cho phần {idx+1}]"

        return "\n".join(texts)

    def convert_to_markdown(self, file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        input_path = file_path

        # Convert .doc sang .docx
        if ext == ".doc":
            docx_path = str(Path(file_path).with_suffix(".docx"))
            try:
                subprocess.run(
                    ["soffice", "--headless", "--convert-to", "docx", "--outdir", str(Path(file_path).parent), file_path],
                    check=True,
                )
                input_path = docx_path
            except subprocess.CalledProcessError:
                raise RuntimeError("Failed to convert .doc to .docx using LibreOffice.")
        elif ext != ".docx":
            raise ValueError("Only .doc or .docx files are supported.")

        logging.info(f"Reading DOCX file: {input_path}")
        full_text = self._extract_text_from_docx(input_path)

        if not full_text.strip():
            raise ValueError("No text content found in the DOCX file.")

        text_chunks = self._split_text(full_text)
        logging.info(f"Document split into {len(text_chunks)} chunk(s)")

        all_md_parts = []

        for idx, chunk in enumerate(tqdm(text_chunks, desc="Converting chunks with LLM")):
            prompt = self._build_prompt(chunk)
            logging.info(f"Sending chunk {idx+1} to LLM...")

            # Retry với backoff khi gặp 503
            for attempt in range(3):
                try:
                    response = self.factory.chat_completion(model_key=self.model_key, messages=prompt)
                    markdown_text = self._extract_response_text(response, idx)
                    all_md_parts.append(markdown_text)
                    logging.info(f"Chunk {idx+1} processed successfully.")
                    break
                except Exception as e:
                    if "503" in str(e) and attempt < 2:
                        wait_time = 2 ** attempt
                        logging.warning(f"Chunk {idx+1} bị 503, thử lại sau {wait_time}s...")
                        time.sleep(wait_time)
                        continue
                    logging.error(f"Error processing chunk {idx+1}: {e}")
                    all_md_parts.append(f"[Lỗi khi xử lý chunk {idx+1}]")
                    break

        # Gộp kết quả
        final_markdown = "\n\n".join(all_md_parts)

        # Lưu file
        safe_name = self._normalize_filename(os.path.basename(input_path))
        md_path = os.path.join(os.path.dirname(input_path), os.path.splitext(safe_name)[0] + "_llm_converted.md")

        with open(md_path, "w", encoding="utf-8") as f:
            f.write(final_markdown)

        logging.info(f"Markdown written to: {md_path}")
        return md_path


if __name__ == "__main__":

    converter = DocxMarkdownConverter()
    file_path = "data/sample_documents/elpasoamiandmdmsrfpv19_10012019-final_described.docx"
    converter.convert_to_markdown(file_path)